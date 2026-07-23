import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from clouda_contracts.security import sanitize_document_text

from .models import PageResult

ARABIC_RE = re.compile(r"[\u0600-\u06FF]")
LATIN_RE = re.compile(r"[A-Za-z]")


def _is_rtl_text(text: str) -> bool:
    arabic = len(ARABIC_RE.findall(text or ""))
    latin = len(LATIN_RE.findall(text or ""))
    return arabic > 0 and arabic >= latin


def _set_paragraph_direction(paragraph, rtl: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    if bidi is not None:
        bidi.set(qn("w:val"), "1" if rtl else "0")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run_rtl = _is_rtl_text(run.text)
        run.font.name = "Arial"
        run._element.get_or_add_rPr().set(qn("w:rtl"), "1" if run_rtl else "0")


def _set_paragraph_rtl(paragraph) -> None:
    _set_paragraph_direction(paragraph, True)


def _add_page_number_field(paragraph, prefix_text: str = "ص ") -> None:
    paragraph.add_run(prefix_text)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def _manual_review_marker(page: PageResult) -> str:
    score = (
        page.text_quality_score
        if page.text_quality_score is not None
        else page.quality_score
    )
    score_text = "unknown" if score is None else f"{float(score):.2f}%"
    reason = (page.review_reason or "estimated quality below 90%").strip()
    model = (page.model_used or "unknown").strip()
    return (
        f"[PAGE {page.page_no} REQUIRES MANUAL REVIEW - "
        f"estimated quality {score_text}; route/model: {model}; reason: {reason}]"
    )


def markdown_to_docx(markdown_pages: list[PageResult]) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        footer_p = (
            section.footer.paragraphs[0]
            if section.footer.paragraphs
            else section.footer.add_paragraph()
        )
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p._p.clear_content()
        _set_paragraph_rtl(footer_p)
        _add_page_number_field(footer_p, prefix_text="ص ")
        if footer_p.runs:
            footer_p.runs[0].font.size = Pt(9)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(12)

    for idx, page in enumerate(markdown_pages):
        if idx > 0:
            doc.add_page_break()

        source_markdown = sanitize_document_text(page.markdown or "")
        if page.requires_manual_review and not source_markdown.lstrip().startswith(
            "[PAGE "
        ):
            source_markdown = f"{_manual_review_marker(page)}\n\n{source_markdown}"
        lines = [ln.rstrip() for ln in source_markdown.splitlines()]
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                p = doc.add_paragraph("")
                _set_paragraph_direction(p, True)
                i += 1
                continue

            if line.startswith("### "):
                p = doc.add_heading(line[4:].strip(), level=3)
                _set_paragraph_direction(p, _is_rtl_text(line))
                i += 1
                continue
            if line.startswith("## "):
                p = doc.add_heading(line[3:].strip(), level=2)
                _set_paragraph_direction(p, _is_rtl_text(line))
                i += 1
                continue
            if line.startswith("# "):
                p = doc.add_heading(line[2:].strip(), level=1)
                _set_paragraph_direction(p, _is_rtl_text(line))
                i += 1
                continue

            if re.match(r"^[-*]\s+", line):
                p = doc.add_paragraph(
                    re.sub(r"^[-*]\s+", "", line), style="List Bullet"
                )
                _set_paragraph_direction(p, _is_rtl_text(line))
                i += 1
                continue
            if re.match(r"^\d+\.\s+", line):
                p = doc.add_paragraph(
                    re.sub(r"^\d+\.\s+", "", line), style="List Number"
                )
                _set_paragraph_direction(p, _is_rtl_text(line))
                i += 1
                continue

            if (
                "|" in line
                and (i + 1) < len(lines)
                and re.match(r"^\|?\s*[-:| ]+\|?\s*$", lines[i + 1].strip())
            ):
                i += 2
                while i < len(lines) and "|" in lines[i]:
                    i += 1
                continue

            p = doc.add_paragraph(line)
            _set_paragraph_direction(p, _is_rtl_text(line))
            i += 1

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
