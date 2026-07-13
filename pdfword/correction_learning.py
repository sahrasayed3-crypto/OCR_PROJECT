import difflib
import hashlib
import io
import json
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

SENSITIVE_PATTERNS = (
    re.compile(r"\b[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}\b"),
    re.compile(r"https?://\S+|www\.\S+", re.I),
    re.compile(r"(?:\+?\d[\s-]?){8,15}"),
    re.compile(r"\b\d{1,4}[/.-]\d{1,2}[/.-]\d{1,4}\b"),
    re.compile(r"\b(?:EGP|USD|EUR|جنيه|دولار|ريال)\s*\d+(?:[.,]\d+)?\b", re.I),
    re.compile(r"\b\d{8,18}\b"),
    re.compile(r"\b[A-Z]{2,}[-_/]?\d{2,}\b"),
)
ARABIC_RE = re.compile(r"[\u0600-\u06ff]")
TOKEN_RE = re.compile(r"\w+|[^\w\s]", re.UNICODE)
SENTENCE_RE = re.compile(r"(?<=[.!?؟؛])\s+|\n+")


def validate_docx_bytes(payload: bytes) -> DocxDocument:
    if len(payload) < 4 or payload[:2] != b"PK":
        raise ValueError("Invalid DOCX content")
    try:
        return Document(io.BytesIO(payload))
    except Exception as exc:
        raise ValueError("Invalid DOCX content") from exc


def file_sha256(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def atomic_save_revision(payload: bytes, directory: Path, version: int) -> Path:
    validate_docx_bytes(payload)
    directory = directory.resolve()
    directory.mkdir(parents=True, exist_ok=True)
    target = (directory / f"corrected_v{version}.docx").resolve()
    target.relative_to(directory)
    with tempfile.NamedTemporaryFile(
        dir=directory, suffix=".docx.part", delete=False
    ) as handle:
        temporary = Path(handle.name)
        handle.write(payload)
        handle.flush()
    temporary.replace(target)
    return target


def _alignment(value) -> str:
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(value, "default")


def _paragraphs(document: DocxDocument) -> list[dict]:
    rows = []
    for index, paragraph in enumerate(document.paragraphs):
        runs = paragraph.runs
        fmt = paragraph.paragraph_format
        rows.append(
            {
                "index": index,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "",
                "bold": any(bool(run.bold) for run in runs),
                "italic": any(bool(run.italic) for run in runs),
                "font_size": next(
                    (round(run.font.size.pt, 2) for run in runs if run.font.size), None
                ),
                "alignment": _alignment(paragraph.alignment),
                "rtl": any(bool(getattr(run.font, "rtl", False)) for run in runs),
                "space_before": fmt.space_before.pt if fmt.space_before else None,
                "space_after": fmt.space_after.pt if fmt.space_after else None,
                "list_type": (
                    "list"
                    if paragraph.style and "list" in paragraph.style.name.lower()
                    else ""
                ),
            }
        )
    return rows


def is_sensitive_text(text: str) -> bool:
    value = (text or "").strip()
    if not value:
        return False
    if value.isdigit():
        return True
    if any(pattern.search(value) for pattern in SENSITIVE_PATTERNS):
        return True
    # Conservative person-name heuristic: short title-cased or Arabic name-like replacements.
    words = value.split()
    if 1 <= len(words) <= 3 and all(
        word[:1].isupper() for word in words if word[:1].isalpha()
    ):
        return True
    return False


def _language(text: str) -> str:
    return "ar" if ARABIC_RE.search(text or "") else "en"


def _change_type(before: str, after: str, opcode: str) -> str:
    if opcode == "insert":
        return "insertion"
    if opcode == "delete":
        return "deletion"
    if before.replace(" ", "") == after.replace(" ", ""):
        return (
            "word_split_merge" if before.count(" ") != after.count(" ") else "spacing"
        )
    if re.sub(r"[^\w\s]", "", before) == re.sub(r"[^\w\s]", "", after):
        return "punctuation"
    return "arabic_character" if _language(before + after) == "ar" else "replacement"


def compare_docx(
    original: str | Path, corrected: str | Path, context_window: int = 30
) -> dict:
    source = _paragraphs(Document(str(original)))
    target = _paragraphs(Document(str(corrected)))
    alignment = difflib.SequenceMatcher(
        a=[row["text"] for row in source],
        b=[row["text"] for row in target],
        autojunk=False,
    )
    examples: list[dict] = []
    formatting: list[dict] = []
    sentence_index = 0
    for tag, a0, a1, b0, b1 in alignment.get_opcodes():
        paired = min(a1 - a0, b1 - b0)
        for offset in range(paired):
            left = source[a0 + offset]
            right = target[b0 + offset]
            for key in (
                "style",
                "bold",
                "italic",
                "font_size",
                "alignment",
                "rtl",
                "space_before",
                "space_after",
                "list_type",
            ):
                if left[key] != right[key]:
                    change = {
                        "paragraph_index": left["index"],
                        "property": key,
                        "before": left[key],
                        "after": right[key],
                    }
                    formatting.append(change)
                    examples.append(
                        {
                            "wrong_text": str(left[key]),
                            "correct_text": str(right[key]),
                            "context_before": left["text"][:context_window],
                            "context_after": "",
                            "paragraph_index": left["index"],
                            "sentence_index": sentence_index,
                            "change_type": f"formatting:{key}",
                            "language": _language(left["text"]),
                            "is_sensitive": False,
                            "status": "proposed",
                        }
                    )
            left_tokens = TOKEN_RE.findall(left["text"])
            right_tokens = TOKEN_RE.findall(right["text"])
            matcher = difflib.SequenceMatcher(
                a=left_tokens, b=right_tokens, autojunk=False
            )
            for op, i1, i2, j1, j2 in matcher.get_opcodes():
                if op == "equal":
                    continue
                before = " ".join(left_tokens[i1:i2])
                after = " ".join(right_tokens[j1:j2])
                context_before = " ".join(left_tokens[max(0, i1 - 8) : i1])[
                    -context_window:
                ]
                context_after = " ".join(left_tokens[i2 : i2 + 8])[:context_window]
                examples.append(
                    {
                        "wrong_text": before,
                        "correct_text": after,
                        "context_before": context_before,
                        "context_after": context_after,
                        "paragraph_index": left["index"],
                        "sentence_index": sentence_index,
                        "change_type": _change_type(before, after, op),
                        "language": _language(before + after),
                        "is_sensitive": is_sensitive_text(before)
                        or is_sensitive_text(after),
                        "status": "proposed",
                    }
                )
            sentence_index += max(1, len(SENTENCE_RE.split(right["text"])))
        for index in range(a0 + paired, a1):
            examples.append(
                {
                    "wrong_text": source[index]["text"],
                    "correct_text": "",
                    "context_before": "",
                    "context_after": "",
                    "paragraph_index": source[index]["index"],
                    "sentence_index": sentence_index,
                    "change_type": "paragraph_deletion",
                    "language": _language(source[index]["text"]),
                    "is_sensitive": is_sensitive_text(source[index]["text"]),
                    "status": "proposed",
                }
            )
        for index in range(b0 + paired, b1):
            examples.append(
                {
                    "wrong_text": "",
                    "correct_text": target[index]["text"],
                    "context_before": "",
                    "context_after": "",
                    "paragraph_index": target[index]["index"],
                    "sentence_index": sentence_index,
                    "change_type": "paragraph_insertion",
                    "language": _language(target[index]["text"]),
                    "is_sensitive": is_sensitive_text(target[index]["text"]),
                    "status": "proposed",
                }
            )
    return {"examples": examples, "formatting_changes": formatting}


def rules_checksum(rules: list[dict]) -> str:
    compact = json.dumps(
        rules, ensure_ascii=False, sort_keys=True, separators=(",", ":")
    )
    return hashlib.sha256(compact.encode("utf-8")).hexdigest()


def dataset_split(job_id: str, evaluation_percent: int = 20) -> str:
    bucket = int(hashlib.sha256(job_id.encode("utf-8")).hexdigest()[:8], 16) % 100
    return "evaluation" if bucket < evaluation_percent else "training"


@dataclass
class RuleApplication:
    text: str
    applications: list[dict]
    conflicts: list[dict]


def apply_correction_rules(
    text: str, rules: list[dict], *, job_id: str = "", document_type: str = ""
) -> RuleApplication:
    eligible = []
    for rule in rules:
        if (
            not rule.get("approved")
            or not rule.get("enabled")
            or rule.get("is_sensitive")
        ):
            continue
        if float(rule.get("confidence") or 0) < float(rule.get("threshold") or 0.9):
            continue
        scope = rule.get("scope", "global")
        if scope == "job_only" and rule.get("scope_value") != job_id:
            continue
        if scope == "document_type" and rule.get("scope_value") != document_type:
            continue
        eligible.append(rule)
    eligible.sort(
        key=lambda row: (-float(row.get("confidence") or 0), int(row.get("id") or 0))
    )
    updated = text
    occupied: set[str] = set()
    applications, conflicts = [], []
    for rule in eligible:
        wrong = str(rule.get("wrong_text") or "")
        correct = str(rule.get("correct_text") or "")
        if wrong in occupied:
            conflicts.append({"rule_id": rule.get("id"), "wrong_text": wrong})
            continue
        if not wrong or wrong not in updated:
            continue
        before_context = str(rule.get("context_pattern") or "")
        if before_context and before_context not in updated:
            continue
        updated, count = updated.replace(wrong, correct), updated.count(wrong)
        if count:
            occupied.add(wrong)
            applications.append(
                {
                    "rule_id": rule.get("id"),
                    "before": wrong,
                    "after": correct,
                    "confidence": float(rule.get("confidence") or 0),
                    "context_match": True,
                    "count": count,
                }
            )
    return RuleApplication(updated, applications, conflicts)
