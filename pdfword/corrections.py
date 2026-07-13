import difflib
from pathlib import Path

from docx import Document

from .database import Database


def extract_docx_text(path: str | Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip())


def propose_corrections(
    original_docx: str | Path, corrected_docx: str | Path, database: Database
) -> dict:
    original_words = extract_docx_text(original_docx).split()
    corrected_words = extract_docx_text(corrected_docx).split()
    matcher = difflib.SequenceMatcher(
        a=original_words, b=corrected_words, autojunk=False
    )
    counts = {"added": 0, "deleted": 0, "replaced": 0}
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        before = " ".join(original_words[i1:i2])
        after = " ".join(corrected_words[j1:j2])
        if (
            tag == "replace"
            and before
            and after
            and len(before) <= 160
            and len(after) <= 160
        ):
            database.add_correction_rule(before, after, "replacement")
            counts["replaced"] += 1
        elif tag == "insert":
            counts["added"] += j2 - j1
        elif tag == "delete":
            counts["deleted"] += i2 - i1
    return counts
