import json
import os
import random
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pypdf import PdfReader, PdfWriter

ROOT = Path(__file__).resolve().parents[1]
SAMPLES_DIR = ROOT / "samples"
GENERATED_DIR = SAMPLES_DIR / "generated"
MANIFEST_PATH = GENERATED_DIR / "cases_manifest.json"
RNG_SEED = 424242
REQUIRED_CATEGORIES = [
    "born_digital_modern",
    "clean_scans",
    "old_archived_scans",
    "low_quality_scans",
    "blurry_docs",
    "fax_style_pdfs",
    "tilted_rotated_scans",
    "mixed_multi_page_docs",
    "encrypted_pdfs",
    "corrupted_pdfs",
]
MIN_SAMPLES_PER_CATEGORY = 2
ARABIC_FIXTURES = {
    "clear_ar": (
        "تحويل ملفات PDF إلى Word بدقة عالية.\n"
        "يحافظ النظام على النص العربي والأرقام وعلامات الترقيم.\n"
        "أجري الاختبار في 5 يوليو 2026.\n"
        "رقم المستند هو 2026-075.\n"
        "المبلغ الإجمالي هو 1250.50 جنيه.\n"
        "هذه عينة عربية موثقة لاختبار دقة استخراج النص."
    ),
    "complex_ar": (
        "تقرير متابعة المشروع\n"
        "بدأ الاجتماع في الساعة 10:30 صباحًا يوم 5 يوليو 2026.\n"
        "المرحلة الأولى: التحقق من النصوص العربية والإنجليزية.\n"
        "المرحلة الثانية: قياس الدقة باستخدام CER و WER.\n"
        "النتيجة المطلوبة لا تقل عن 90% للعينات المقبولة.\n"
        "Clouda PDF يحول المستند إلى Word مع الحفاظ على ترتيب الفقرات."
    ),
}


def _render_pages(
    pdf_path: Path, dpi: int = 260, max_pages: int = 3
) -> list[Image.Image]:
    pages: list[Image.Image] = []
    doc = pdfium.PdfDocument(str(pdf_path))
    try:
        total = min(len(doc), max_pages)
        for i in range(total):
            page = doc.get_page(i)
            bitmap = None
            try:
                bitmap = page.render(scale=dpi / 72.0)
                pil = bitmap.to_pil().convert("RGB")
                pages.append(pil.copy())
                pil.close()
            finally:
                if bitmap is not None and hasattr(bitmap, "close"):
                    bitmap.close()
                page.close()
    finally:
        doc.close()
    return pages


def _save_pdf(images: list[Image.Image], out_path: Path) -> None:
    if not images:
        raise ValueError("no images to save")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    first = images[0].convert("RGB")
    rest = [img.convert("RGB") for img in images[1:]]
    first.save(out_path, "PDF", resolution=300.0, save_all=True, append_images=rest)
    first.close()
    for img in rest:
        img.close()
    _normalize_pdf(out_path)


def _normalize_pdf(path: Path) -> None:
    """Remove generator timestamps and volatile IDs from committed fixtures."""
    reader = PdfReader(path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": "Clouda PDF deterministic benchmark fixture",
            "/Creator": "Clouda fixture generator",
            "/Producer": "pypdf",
            "/CreationDate": "D:20260705000000+00'00'",
            "/ModDate": "D:20260705000000+00'00'",
        }
    )
    temporary = path.with_suffix(".normalized.pdf")
    with temporary.open("wb") as output:
        writer.write(output)
    temporary.replace(path)


def _clean_scan(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    gray = ImageOps.autocontrast(gray, cutoff=1)
    return gray.convert("RGB")


def _deterministic_noise(size: tuple[int, int], seed: int) -> Image.Image:
    rnd = random.Random(seed)
    w, h = size
    data = bytes(rnd.randint(0, 255) for _ in range(w * h))
    return Image.frombytes("L", size, data)


def _old_scan(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    low_contrast = ImageEnhance.Contrast(gray).enhance(0.72)
    blur = low_contrast.filter(ImageFilter.GaussianBlur(radius=1.0))
    base = ImageOps.colorize(blur, black="#4b3c2d", white="#f2e7ca")
    noise = _deterministic_noise(
        base.size, seed=RNG_SEED + (base.size[0] * 31) + base.size[1]
    )
    noise = ImageEnhance.Contrast(noise).enhance(1.4)
    noise_rgb = Image.merge("RGB", (noise, noise, noise))
    mixed = Image.blend(base, noise_rgb, alpha=0.10)
    return mixed


def _low_contrast(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    soft = ImageEnhance.Contrast(gray).enhance(0.55)
    soft = ImageEnhance.Brightness(soft).enhance(1.08)
    return soft.convert("RGB")


def _blurry_scan(img: Image.Image) -> Image.Image:
    w, h = img.size
    small = img.resize((max(600, w // 2), max(800, h // 2)), Image.Resampling.BILINEAR)
    back = small.resize((w, h), Image.Resampling.BILINEAR)
    return back.filter(ImageFilter.GaussianBlur(radius=1.4))


def _fax_like(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    gray = ImageEnhance.Contrast(gray).enhance(1.45)
    bw = gray.point(lambda p: 255 if p > 165 else 0, mode="1")
    return bw.convert("RGB")


def _tilted_scan(img: Image.Image) -> Image.Image:
    rotated = img.rotate(1.4, expand=True, fillcolor=(250, 250, 250))
    return rotated.filter(ImageFilter.UnsharpMask(radius=1.2, percent=150, threshold=2))


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _copy_reference(src_ref: Path, out_ref: Path) -> str:
    text = src_ref.read_text(encoding="utf-8")
    _write_text(out_ref, text)
    return text


def _set_arabic_run_font(run, font_name: str = "Arial") -> None:
    run.font.name = font_name
    run.font.size = Pt(18)
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:cs"), font_name)


def _create_arabic_docx(path: Path, text: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    for line in text.splitlines():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_properties = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        paragraph_properties.append(bidi)
        _set_arabic_run_font(paragraph.add_run(line))
    document.save(str(path))


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Create a selectable Unicode text layer without redistributing a font."""
    if shutil.which("powershell") and Path(r"C:\Windows\Fonts\arial.ttf").is_file():
        script = (
            "$word=New-Object -ComObject Word.Application;"
            "$word.Visible=$false;"
            "try{$doc=$word.Documents.Open($env:CLOUD_FIXTURE_DOCX);"
            "$doc.SaveAs([ref]$env:CLOUD_FIXTURE_PDF,[ref]17);$doc.Close()}"
            "finally{$word.Quit()}"
        )
        env = dict(os.environ)
        env["CLOUD_FIXTURE_DOCX"] = str(docx_path.resolve())
        env["CLOUD_FIXTURE_PDF"] = str(pdf_path.resolve())
        subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
            check=True,
            env=env,
            timeout=60,
        )
        return
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if office:
        subprocess.run(
            [
                office,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            check=True,
            timeout=60,
        )
        generated = pdf_path.parent / f"{docx_path.stem}.pdf"
        if generated != pdf_path:
            generated.replace(pdf_path)
        return
    raise RuntimeError(
        "Arabic fixture generation requires Microsoft Word or LibreOffice"
    )


def _create_arabic_sources() -> None:
    for key, text in ARABIC_FIXTURES.items():
        reference_path = SAMPLES_DIR / f"sample_{key}_ref.txt"
        pdf_path = SAMPLES_DIR / f"sample_{key}.pdf"
        _write_text(reference_path, text + "\n")
        with tempfile.TemporaryDirectory(prefix="clouda-fixture-") as temp_name:
            docx_path = Path(temp_name) / f"sample_{key}.docx"
            temporary_pdf = Path(temp_name) / f"sample_{key}.pdf"
            _create_arabic_docx(docx_path, text)
            _convert_docx_to_pdf(docx_path, temporary_pdf)
            replacement = pdf_path.with_suffix(".fixture-new.pdf")
            shutil.copyfile(temporary_pdf, replacement)
            for attempt in range(5):
                try:
                    replacement.replace(pdf_path)
                    break
                except OSError:
                    if attempt == 4:
                        raise
                    time.sleep(0.2)
            _normalize_pdf(pdf_path)


def _add_case(
    cases: list[dict],
    case_id: str,
    pdf_path: Path,
    ref_path: Path | None,
    language: str,
    category: str,
    expect_failure: bool = False,
) -> None:
    entry = {
        "id": case_id,
        "pdf": str(pdf_path.relative_to(ROOT)).replace("\\", "/"),
        "language": language,
        "category": category,
        "pages": [1],
        "expect_failure": expect_failure,
    }
    if ref_path is not None:
        entry["reference"] = str(ref_path.relative_to(ROOT)).replace("\\", "/")
    cases.append(entry)


def _category_counts(cases: list[dict]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for c in cases:
        cat = str(c.get("category", "")).strip()
        counts[cat] = counts.get(cat, 0) + 1
    return counts


def validate_manifest_cases(cases: list[dict]) -> list[str]:
    errors: list[str] = []
    counts = _category_counts(cases)
    for cat in REQUIRED_CATEGORIES:
        count = counts.get(cat, 0)
        if count == 0:
            errors.append(f"missing category: {cat}")
        elif count < MIN_SAMPLES_PER_CATEGORY:
            errors.append(f"category has too few samples ({count}): {cat}")
    return errors


def main() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    _create_arabic_sources()
    cases: list[dict] = []

    bases = [
        (
            "clear_ar",
            SAMPLES_DIR / "sample_clear_ar.pdf",
            SAMPLES_DIR / "sample_clear_ar_ref.txt",
            "ar",
        ),
        (
            "complex_ar",
            SAMPLES_DIR / "sample_complex_ar.pdf",
            SAMPLES_DIR / "sample_complex_ar_ref.txt",
            "ar",
        ),
        (
            "clear_en",
            SAMPLES_DIR / "sample_clear_en.pdf",
            SAMPLES_DIR / "sample_clear_en_ref.txt",
            "en",
        ),
        (
            "complex_en",
            SAMPLES_DIR / "sample_complex_en.pdf",
            SAMPLES_DIR / "sample_complex_en_ref.txt",
            "en",
        ),
    ]

    transforms = [
        ("clean_scans", _clean_scan),
        ("old_archived_scans", _old_scan),
        ("low_quality_scans", _low_contrast),
        ("blurry_docs", _blurry_scan),
        ("fax_style_pdfs", _fax_like),
        ("tilted_rotated_scans", _tilted_scan),
    ]

    for key, pdf_path, ref_path, language in bases:
        if not pdf_path.exists() or not ref_path.exists():
            continue

        # Original born-digital case.
        _add_case(
            cases, f"{key}_modern", pdf_path, ref_path, language, "born_digital_modern"
        )

        if language == "en":
            for category_name, _tf_fn in transforms:
                out_pdf = GENERATED_DIR / f"{key}_{category_name}.pdf"
                out_ref = GENERATED_DIR / f"{key}_{category_name}_ref.txt"
                _add_case(
                    cases,
                    f"{key}_{category_name}",
                    out_pdf,
                    out_ref,
                    language,
                    category_name,
                )
            continue

        page_images = _render_pages(pdf_path, dpi=320, max_pages=2)

        for category_name, tf_fn in transforms:
            transformed = [tf_fn(img) for img in page_images]
            out_pdf = GENERATED_DIR / f"{key}_{category_name}.pdf"
            out_ref = GENERATED_DIR / f"{key}_{category_name}_ref.txt"
            _save_pdf(transformed, out_pdf)
            _copy_reference(ref_path, out_ref)
            _add_case(
                cases,
                f"{key}_{category_name}",
                out_pdf,
                out_ref,
                language,
                category_name,
            )
            for img in transformed:
                img.close()

        for img in page_images:
            img.close()

    # Mixed multipage Arabic file (modern + old scan style).
    clear_ar_pdf = SAMPLES_DIR / "sample_clear_ar.pdf"
    complex_ar_pdf = SAMPLES_DIR / "sample_complex_ar.pdf"
    clear_ar_ref = SAMPLES_DIR / "sample_clear_ar_ref.txt"
    complex_ar_ref = SAMPLES_DIR / "sample_complex_ar_ref.txt"
    if clear_ar_pdf.exists() and complex_ar_pdf.exists():
        mix_pages = _render_pages(clear_ar_pdf, dpi=280, max_pages=1) + _render_pages(
            complex_ar_pdf, dpi=300, max_pages=1
        )
        mixed = [_clean_scan(mix_pages[0]), _old_scan(mix_pages[1])]
        multi_pdf = GENERATED_DIR / "arabic_mixed_multipage.pdf"
        _save_pdf(mixed, multi_pdf)
        for img in mix_pages + mixed:
            img.close()

        mix_ref_txt = ""
        if clear_ar_ref.exists():
            mix_ref_txt += clear_ar_ref.read_text(encoding="utf-8").strip()
        if complex_ar_ref.exists():
            mix_ref_txt += "\n\n" + complex_ar_ref.read_text(encoding="utf-8").strip()
        mix_ref = GENERATED_DIR / "arabic_mixed_multipage_ref.txt"
        _write_text(mix_ref, mix_ref_txt.strip())
        cases.append(
            {
                "id": "arabic_mixed_multipage",
                "pdf": str(multi_pdf.relative_to(ROOT)).replace("\\", "/"),
                "reference": str(mix_ref.relative_to(ROOT)).replace("\\", "/"),
                "language": "ar",
                "category": "mixed_multi_page_docs",
                "pages": [1, 2],
                "expect_failure": False,
            }
        )
    clear_en_pdf = SAMPLES_DIR / "sample_clear_en.pdf"
    complex_en_pdf = SAMPLES_DIR / "sample_complex_en.pdf"
    if clear_en_pdf.exists() and complex_en_pdf.exists():
        multi_pdf = GENERATED_DIR / "english_mixed_multipage.pdf"
        mix_ref = GENERATED_DIR / "english_mixed_multipage_ref.txt"
        cases.append(
            {
                "id": "english_mixed_multipage",
                "pdf": str(multi_pdf.relative_to(ROOT)).replace("\\", "/"),
                "reference": str(mix_ref.relative_to(ROOT)).replace("\\", "/"),
                "language": "en",
                "category": "mixed_multi_page_docs",
                "pages": [1, 2],
                "expect_failure": False,
            }
        )

    # Encrypted PDF expected failure.
    encrypted_pdf = GENERATED_DIR / "encrypted_sample.pdf"
    encrypted_pdf_en = GENERATED_DIR / "encrypted_sample_en.pdf"
    if clear_ar_pdf.exists() and not encrypted_pdf.exists():
        reader = PdfReader(str(clear_ar_pdf))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt("1234")
        with encrypted_pdf.open("wb") as f:
            writer.write(f)
    if encrypted_pdf.exists():
        _add_case(
            cases,
            "encrypted_pdf_expected_fail_ar",
            encrypted_pdf,
            None,
            "ar",
            "encrypted_pdfs",
            expect_failure=True,
        )
    if clear_en_pdf.exists() and not encrypted_pdf_en.exists():
        reader = PdfReader(str(clear_en_pdf))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt("1234")
        with encrypted_pdf_en.open("wb") as f:
            writer.write(f)
    if encrypted_pdf_en.exists():
        _add_case(
            cases,
            "encrypted_pdf_expected_fail_en",
            encrypted_pdf_en,
            None,
            "en",
            "encrypted_pdfs",
            expect_failure=True,
        )

    # Corrupted PDF expected failure.
    corrupted_pdf = GENERATED_DIR / "corrupted_sample.pdf"
    corrupted_pdf_en = GENERATED_DIR / "corrupted_sample_en.pdf"
    if clear_ar_pdf.exists() and not corrupted_pdf.exists():
        raw = clear_ar_pdf.read_bytes()
        corrupted_pdf.write_bytes(raw[: max(100, len(raw) // 5)])
    if corrupted_pdf.exists():
        _add_case(
            cases,
            "corrupted_pdf_expected_fail_ar",
            corrupted_pdf,
            None,
            "ar",
            "corrupted_pdfs",
            expect_failure=True,
        )
    if clear_en_pdf.exists() and not corrupted_pdf_en.exists():
        raw = clear_en_pdf.read_bytes()
        corrupted_pdf_en.write_bytes(raw[: max(100, len(raw) // 5)])
    if corrupted_pdf_en.exists():
        _add_case(
            cases,
            "corrupted_pdf_expected_fail_en",
            corrupted_pdf_en,
            None,
            "en",
            "corrupted_pdfs",
            expect_failure=True,
        )

    errors = validate_manifest_cases(cases)
    if errors:
        raise RuntimeError("Invalid benchmark manifest: " + "; ".join(errors))

    payload = {
        "fixture_version": 2,
        "cases": cases,
    }
    MANIFEST_PATH.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"Generated {len(cases)} cases.")
    print(f"Manifest: {MANIFEST_PATH}")


if __name__ == "__main__":
    main()
