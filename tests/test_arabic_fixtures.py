import io
import json
import re
from pathlib import Path

import fitz
import pdfplumber
from docx import Document
from pypdf import PdfReader

from pdfword.accuracy import compute_accuracy_metrics
from pdfword.docx_export import markdown_to_docx
from pdfword.models import PageResult
from pdfword.ocr_pipeline import _clean_markdown_output, process_pdf

ROOT = Path(__file__).resolve().parents[1]
MANIFEST = ROOT / "samples" / "generated" / "cases_manifest.json"
ARABIC_RE = re.compile(r"[\u0600-\u06ff]")


def _arabic_reference_paths() -> list[Path]:
    payload = json.loads(MANIFEST.read_text(encoding="utf-8"))
    return sorted(
        {
            ROOT / case["reference"]
            for case in payload["cases"]
            if case.get("language") == "ar" and case.get("reference")
        }
    )


def test_arabic_references_are_strict_utf8_unicode() -> None:
    paths = _arabic_reference_paths()
    assert paths
    for path in paths:
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="strict")
        assert text.encode("utf-8").decode("utf-8") == text
        assert ARABIC_RE.search(text), path
        assert "???" not in text, path
        assert text.count("?") <= 1, path


def test_born_digital_arabic_has_real_text_layer_in_three_extractors() -> None:
    pdf_path = ROOT / "samples" / "sample_complex_ar.pdf"
    reference = (ROOT / "samples" / "sample_complex_ar_ref.txt").read_text(
        encoding="utf-8"
    )
    data = pdf_path.read_bytes()

    with fitz.open(stream=data, filetype="pdf") as document:
        pymupdf_text = "\n".join(page.get_text("text") for page in document)
    with pdfplumber.open(io.BytesIO(data)) as document:
        pdfplumber_text = "\n".join(
            (page.extract_text() or "") for page in document.pages
        )
    pypdf_text = "\n".join(
        (page.extract_text() or "") for page in PdfReader(io.BytesIO(data)).pages
    )

    for extracted in (pymupdf_text, pdfplumber_text, pypdf_text):
        assert ARABIC_RE.search(extracted)
        assert len(extracted.strip()) >= 100

    metrics = compute_accuracy_metrics(reference, pypdf_text)
    assert metrics["cer"] <= 10.0
    # pypdf preserves the Arabic text layer but may reorder digits/punctuation in
    # mixed Arabic-English lines, so character accuracy is the stable readiness
    # signal here rather than an OCR-quality claim.
    assert metrics["character_accuracy"] >= 90.0


def test_clean_markdown_removes_duplicate_lines_without_ocr_repairs() -> None:
    assert _clean_markdown_output("hello\nhello\nworld") == "hello\nworld"


def test_scanned_arabic_fixtures_wait_for_future_ocr_engine() -> None:
    for filename in ("clear_ar_clean_scans.pdf", "clear_ar_low_quality_scans.pdf"):
        path = ROOT / "samples" / "generated" / filename
        reader = PdfReader(path)
        assert not "".join((page.extract_text() or "") for page in reader.pages).strip()
        rows, text = process_pdf(
            path.read_bytes(),
            None,
            None,
            page_numbers=[1],
            progress_bar=None,
            status_placeholder=None,
        )
        assert rows[0].route_used == "pending_ocr_model"
        assert rows[0].requires_manual_review is True
        assert "future OCR model" in text


def test_text_only_docx_opens_without_images_or_tables() -> None:
    content = (
        "عنوان عربي\n\nفقرة قابلة للتحرير.\n\n"
        "| عمود أول | عمود ثان |\n|---|---|\n| بيانات | يجب تجاهلها |"
    )
    payload = markdown_to_docx(
        [PageResult(page_no=1, model_used="test", markdown=content)]
    )
    document = Document(io.BytesIO(payload))
    assert "عنوان عربي" in "\n".join(p.text for p in document.paragraphs)
    assert not document.tables
    assert not document.inline_shapes
